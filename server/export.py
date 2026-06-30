from pathlib import Path

from docx import Document
from docx.shared import Pt

from server.backends.base import TranscriptSegment
from server.silence import PAUSE_SPEAKER


def _fmt_ts(seconds: float) -> str:
    total = int(seconds)
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def export_to_docx(
    segments: list[TranscriptSegment],
    output_path: Path,
    *,
    title: str = "Расшифровка судебного заседания",
    speaker_names: dict[str, str] | None = None,
    show_speakers: bool = True,
) -> Path:
    speaker_names = speaker_names or {}
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading(title, level=1)

    last_speaker: str | None = None
    paragraph = None
    for seg in segments:
        is_pause = seg.speaker == PAUSE_SPEAKER

        if is_pause:
            # Пауза — отдельный курсивный абзац, без префикса спикера
            p = doc.add_paragraph()
            run = p.add_run(f"[{_fmt_ts(seg.start)}] {seg.text}")
            run.italic = True
            last_speaker = None
            paragraph = None
            continue

        if not show_speakers:
            # Режим "только текст" — без префикса, без склеивания по спикеру
            p = doc.add_paragraph()
            ts_run = p.add_run(f"[{_fmt_ts(seg.start)}] ")
            ts_run.bold = True
            p.add_run(seg.text)
            last_speaker = None
            paragraph = None
            continue

        display_speaker = speaker_names.get(seg.speaker, seg.speaker)
        if display_speaker != last_speaker:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"[{_fmt_ts(seg.start)}] {display_speaker}: ")
            run.bold = True
            paragraph.add_run(seg.text)
            last_speaker = display_speaker
        else:
            assert paragraph is not None
            paragraph.add_run(" " + seg.text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
